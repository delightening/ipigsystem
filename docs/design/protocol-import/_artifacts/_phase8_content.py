# -*- coding: utf-8 -*-
"""階段8：補抽 _phase4 漏掉的三個「源頭有資料」欄位（跨自由文字版與標準 AD-04-01-01 版）：
  purpose.replacement.alt_search.{platforms, keywords, conclusion}  ← 2.2.2 替代搜尋（標準表 T6）
  purpose.duplicate.status                                          ← 2.2.3 重複性（標準表 T7）
  design.endpoints.{experimental_endpoint, humane_endpoint}         ← 4.1.7（兩版皆把人道終點併在同格）
產出 protocol-content-enrich-fix.json 餵 enrich_imported_protocols（deep_merge，僅 import_pending）。純讀取。"""
import os, re, json, unicodedata
import docx

OUT = r'C:\System Coding\ipig_system\docs\design\protocol-content-enrich-fix.json'
import _phase6_basic as P   # 重用 find_folder / latest_application

PLATFORMS = [
    ('altbib', ['ALTBIB']),
    ('db_alm', ['DB-ALM', 'DBALM']),
    ('re_place', ['歐洲動物替代試驗資源平台', 'RE-PLACE', 'REPLACE']),
    ('johns_hopkins', ['JOHNS HOPKINS', '約翰霍普金斯', 'CAAT']),
    ('taat', ['臺灣', '台灣', 'TAAT']),
    ('nc3rs_eda', ['NC3RS EDA', 'EDA']),
    ('nc3rs_refinement', ['NC3RS 精緻', 'NC3RS REFINE']),
]
UNCHECKED = '□☐'


def nf(s):
    return unicodedata.normalize('NFKC', s or '').strip()


def iter_cells(doc):
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                yield nf(c.text.replace('\n', ' '))


def iter_rows(doc):
    """逐列合併儲存格文字（否/是 等跨格選項用）。"""
    for t in doc.tables:
        for row in t.rows:
            cells = [nf(c.text.replace('\n', ' ')) for c in row.cells]
            out = []
            for c in cells:
                if not out or out[-1] != c:
                    out.append(c)
            yield ' '.join(out)


def parse_alt_search(text):
    """T6：'已於■ALTBIB ■DB-ALM ■歐洲... □其他 搜尋關鍵字(...)並確認...'。"""
    up = text.upper()
    plats = []
    for key, names in PLATFORMS:
        for nm in names:
            i = up.find(nm.upper())
            if i < 0:
                continue
            # 往前找最近的標記字元（非空白）
            j = i - 1
            while j >= 0 and text[j] in ' 　':
                j -= 1
            mark = text[j] if j >= 0 else ''
            if mark not in UNCHECKED:                 # ■ 或無框 = 選取
                plats.append(key)
            break
    if '其他' in text:
        m = re.search(r'其他[:：]?\s*([^\s）)。]+)', text)
        if m and m.group(1) and m.group(1) not in ('___________', '_'):
            # 其他平台有填名稱才算
            if '■其他' in text or re.search(r'[^□☐]\s*其他', text):
                pass
    kw = None
    m = re.search(r'關鍵字\s*[（(]([^）)]+)[）)]', text)
    if m:
        kw = m.group(1).strip()
    concl = None
    m = re.search(r'(並?確認[^。]*。?|結論[:：][^。]*。?)', text)
    if m:
        concl = m.group(1).strip()
    if not (plats or kw or concl):
        return None
    out = {}
    if plats:
        out['platforms'] = plats
    if kw:
        out['keywords'] = kw
    if concl:
        out['conclusion'] = concl
    return out


def parse_duplicate(text):
    """T7：'■否(No) □是(Yes) 請述明重複實驗...'。回 status。
    僅認真正 radio 列（含雙語標記 (No)/(Yes)），避免散文含「重複/是」誤觸。"""
    up = text.upper()
    if '(NO)' not in up or '(YES' not in up:
        return None
    if '重複' not in text:
        return None
    # 找「否」「是」的勾選狀態
    def checked(label):
        i = text.find(label)
        if i < 0:
            return False
        j = i - 1
        while j >= 0 and text[j] in ' 　':
            j -= 1
        return j >= 0 and text[j] not in UNCHECKED
    yes_just = None
    if checked('是'):
        return ('yes_duplicate', None)
    if checked('否'):
        return ('no', None)
    return None


def parse_endpoints(text):
    """同時含「實驗終點」+「人道終點」的終點格 → (experimental, humane)。
    （僅含「人道終點」一詞者多為 2.2.1 散文誤觸，排除。）"""
    if '人道終點' not in text or '實驗終點' not in text:
        return None
    idx = text.find('人道終點')
    exp = text[:idx]
    hum = text[idx + len('人道終點'):]
    exp = re.sub(r'^\s*實驗終點\s*[:：]?\s*[（(]?請說明[）)]?\s*[:：]?', '', exp).strip()
    hum = re.sub(r'^\s*[:：]?\s*', '', hum).strip()
    out = {}
    if exp:
        out['experimental_endpoint'] = exp
    if hum:
        out['humane_endpoint'] = hum
    return out if out else None


def extract(pig):
    fp = P.find_folder(pig)
    if not fp:
        return None, 'no folder'
    path = P.latest_application(fp)
    if not path:
        return None, 'no docx'
    try:
        doc = docx.Document(path)
    except Exception as e:  # noqa: BLE001
        return None, 'docx err: %s' % e

    alt, dup, eps = None, None, None
    for txt in iter_cells(doc):
        if alt is None and ('替代' in txt or '關鍵字' in txt) and ('ALTBIB' in txt.upper() or '關鍵字' in txt):
            alt = parse_alt_search(txt)
        if eps is None:
            e = parse_endpoints(txt)
            if e:
                eps = e
    for txt in iter_rows(doc):       # 否/是 跨格 → 用列合併
        if dup is None:
            d = parse_duplicate(txt)
            if d:
                dup = d

    purpose, design = {}, {}
    if alt:
        purpose.setdefault('replacement', {})['alt_search'] = alt
    if dup:
        purpose['duplicate'] = {'status': dup[0]}
    if eps:
        design['endpoints'] = eps
    overlay = {}
    if purpose:
        overlay['purpose'] = purpose
    if design:
        overlay['design'] = design
    if not overlay:
        return None, 'nothing'
    diag = 'alt:%s dup:%s endpoints:%s' % (
        ('P%d+kw%s+c%s' % (len(alt.get('platforms', [])), 'Y' if alt.get('keywords') else '-',
                           'Y' if alt.get('conclusion') else '-')) if alt else '-',
        dup[0] if dup else '-',
        ('exp%s+hum%s' % ('Y' if eps.get('experimental_endpoint') else '-',
                          'Y' if eps.get('humane_endpoint') else '-')) if eps else '-')
    return overlay, diag


if __name__ == '__main__':
    import sys
    args = sys.argv[1:]
    pigs = P.PIGS if (not args or args == ['--all']) else args
    write = (not args or args == ['--all'])
    result = {}
    for pig in pigs:
        overlay, diag = extract(pig)
        if overlay is None:
            print('SKIP %-12s | %s' % (pig, diag))
            continue
        result[pig] = overlay
        print('OK   %-12s | %s' % (pig, diag))
    if write:
        json.dump(result, open(OUT, 'w', encoding='utf-8'), ensure_ascii=False, indent=1)
        print('\n寫出 %d 筆 → %s' % (len(result), OUT))
    else:
        print('\n（樣本模式，未寫檔）')
